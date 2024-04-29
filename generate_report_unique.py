import pandas as pd
import matplotlib.pyplot as plt
import requests
from PIL import Image
from io import BytesIO
from matplotlib.offsetbox import OffsetImage, AnnotationBbox
from docx import Document
from docx.shared import Inches

# Read the Excel file
df = pd.read_excel('data_with_images.xlsx')

# Group the data by the 'who' column
grouped = df.groupby('who')
count = 0

# Create a Word document
doc = Document()

# Iterate over each unique URL in the 'who' column
for url, group_df in grouped:
    count += 1
    print("current outlet: ", count)

    # Extract headlines, polarity, and subjectivity into a list of tuples
    data = [(row['headline'], row['polarity'], row['subjectivity']) for _, row in group_df.iterrows()]

    # Calculate the average polarity and subjectivity
    avg_polarity = sum(polarity for _, polarity, _ in data) / len(data)
    avg_subjectivity = sum(subjectivity for _, _, subjectivity in data) / len(data)

    # Calculate the average subjectivity for positive and negative polarity scores separately
    positive_data = [(headline, polarity, subjectivity) for headline, polarity, subjectivity in data if polarity > 0]
    negative_data = [(headline, polarity, subjectivity) for headline, polarity, subjectivity in data if polarity < 0]

    avg_subjectivity_positive = sum(subjectivity for _, _, subjectivity in positive_data) / len(positive_data) if len(positive_data) > 0 else 0
    avg_subjectivity_negative = sum(subjectivity for _, _, subjectivity in negative_data) / len(negative_data) if len(negative_data) > 0 else 0

    # Fetch the image using the URL and convert it to PNG format
    response = requests.get(url)
    image = Image.open(BytesIO(response.content))

    # Increase the resolution of the image
    image_width, image_height = image.size
    new_image = image.resize((image_width * 2, image_height * 2))

    # Save the image to a file
    image_filename = f"outlet_logos\\{count}_logo.png"
    new_image.save(image_filename)

    # Add the logo to the Word document
    doc.add_picture(image_filename, width=Inches(2))
    doc.add_paragraph()

    # Add the list of headlines to the Word document as bullet points
    doc.add_heading('Headlines:', level=2)
    for headline, _, _ in data:
        doc.add_paragraph(headline, style='List Bullet')

    # Create a colormap based on 'subjectivity'
    cmap = plt.get_cmap('coolwarm')

    # Plot the data with colored dots
    fig, ax = plt.subplots(figsize=(10, 6))
    scatter = ax.scatter(range(len(data)), [polarity for _, polarity, _ in data],
                          c=[subjectivity for _, _, subjectivity in data], cmap=cmap)

    # Display the image on the graph
    imagebox = OffsetImage(new_image, zoom=0.5)  # Use 0.2 instead of 0.4 for 2x resolution
    ab = AnnotationBbox(imagebox, (0.5, 0.95), xycoords='axes fraction', frameon=False, pad=0.5)
    ax.add_artist(ab)

    plt.colorbar(plt.cm.ScalarMappable(norm=None, cmap='coolwarm'), ax=ax, label='Subjectivity')
    ax.set_xlabel('Headline Index')
    ax.set_ylabel('Polarity')
    ax.set_title(f'Analysis Results')

    plt.subplots_adjust(top=0.8)  # Adjust the top margin to make room for the image

    # Save the plot to a file
    plot_filename = f"outlet_logos\\{count}_plot.png"
    plt.savefig(plot_filename)

    # Add the plot to the Word document
    doc.add_picture(plot_filename, width=Inches(6))
    doc.add_paragraph()

    # Add the average polarity and subjectivity to the Word document
    doc.add_paragraph(f"Average Polarity: {avg_polarity:.2f}")
    doc.add_paragraph(f"Average Subjectivity: {avg_subjectivity:.2f}")
    doc.add_paragraph(f"Average Subjectivity for Positive Polarity: {avg_subjectivity_positive:.2f}")
    doc.add_paragraph(f"Average Subjectivity for Negative Polarity: {avg_subjectivity_negative:.2f}")

    plt.close()
    doc.add_page_break()


doc.save('analysis_results.docx')
