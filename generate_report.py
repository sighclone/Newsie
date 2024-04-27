import openpyxl
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import requests
from io import BytesIO
from PIL import Image

# Load the Excel file
wb = openpyxl.load_workbook("data_with_images.xlsx")
sheet = wb.active

# Create a dictionary to store image URLs and their associated headlines
image_headlines = {}
for row in sheet.iter_rows(min_row=2, values_only=True):
    image_url, headline, _ = row
    if image_url not in image_headlines:
        image_headlines[image_url] = []
    image_headlines[image_url].append(headline)

# Create a new Word document
doc = Document()
doc.add_heading('Images and Headlines', level=1)

# Add each unique image and its associated headlines to the document
for image_url, headlines in image_headlines.items():
    response = requests.get(image_url)
    if response.status_code == 200:
        image = Image.open(BytesIO(response.content))
        # Convert image to PNG
        png_image = BytesIO()
        image.save(png_image, format='PNG')
        png_image.seek(0)
        doc.add_picture(png_image, width=Inches(3))
    bold_text = doc.add_paragraph()
    run = bold_text.add_run("Associated headlines:")
    run.bold = True
    doc.add_paragraph()
    for headline in headlines:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(headline)
    doc.add_paragraph("--------------------------------------------------")

# Save the Word document
doc.save("gathered.docx")

print("Word document 'gathered.docx' has been created.")
